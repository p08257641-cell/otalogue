import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml

doc = Document()

# Set standard page margins (0.9 inch)
for section in doc.sections:
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

# Colors
COLOR_PRIMARY = RGBColor(13, 118, 110)      # Deep Forest Teal #0D766E
COLOR_DARK_GREEN = RGBColor(4, 47, 46)     # Dark Forest Green #042F2E
COLOR_GOLD = RGBColor(217, 119, 6)         # Gold / Amber #D97706
COLOR_DARK = RGBColor(15, 23, 42)          # Slate Dark #0F172A
COLOR_MUTED = RGBColor(71, 85, 105)        # Muted Slate #475569

# 1. Organization Header / Letterhead
p_org = doc.add_paragraph()
p_org.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_org.paragraph_format.space_after = Pt(2)

run_org1 = p_org.add_run('FAMILY HEALTH UNIVERSITY MEDICAL SCHOOL\n')
run_org1.font.name = 'Georgia'
run_org1.font.size = Pt(13)
run_org1.font.bold = True
run_org1.font.color.rgb = COLOR_DARK_GREEN

run_org2 = p_org.add_run('MEDICAL STUDENTS\' ASSOCIATION (FHUMSA)\n')
run_org2.font.name = 'Georgia'
run_org2.font.size = Pt(11)
run_org2.font.bold = True
run_org2.font.color.rgb = COLOR_PRIMARY

run_sub = p_org.add_run('OFFICE OF THE FINANCIAL SECRETARY — ELECT\n')
run_sub.font.name = 'Arial'
run_sub.font.size = Pt(9.5)
run_sub.font.bold = True
run_sub.font.color.rgb = COLOR_GOLD

run_address = p_org.add_run('Family Health University College · P.O. Box TS 669, Teshie, Accra, Ghana\n')
run_address.font.name = 'Arial'
run_address.font.size = Pt(8.5)
run_address.font.color.rgb = COLOR_MUTED

# Divider Rule
p_div = doc.add_paragraph()
p_div.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_div.paragraph_format.space_after = Pt(10)
p_div_border = parse_xml(r'<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:bottom w:val="single" w:sz="18" w:space="1" w:color="0D766E"/></w:pBdr>')
p_div._p.get_or_add_pPr().append(p_div_border)

# 2. Metadata Box (Date, Release Status)
p_meta = doc.add_paragraph()
p_meta.paragraph_format.space_after = Pt(10)

run_for_release = p_meta.add_run('FOR IMMEDIATE RELEASE\n')
run_for_release.font.name = 'Arial'
run_for_release.font.size = Pt(9.5)
run_for_release.font.bold = True
run_for_release.font.color.rgb = COLOR_GOLD

run_date = p_meta.add_run('DATE: August 2026\n')
run_date.font.name = 'Arial'
run_date.font.size = Pt(9)
run_date.font.bold = True
run_date.font.color.rgb = COLOR_DARK

run_contact = p_meta.add_run('SUBJECT: Official Post-Election Appreciation & Acceptance Statement\nISSUED BY: Charlene Odei Asare, Financial Secretary — Elect\n')
run_contact.font.name = 'Arial'
run_contact.font.size = Pt(9)
run_contact.font.color.rgb = COLOR_MUTED

# 3. Main Headline
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title.paragraph_format.space_before = Pt(4)
p_title.paragraph_format.space_after = Pt(14)

run_title = p_title.add_run('WE DID IT: A MANDATE FOR SERVICE & IMPACT')
run_title.font.name = 'Georgia'
run_title.font.size = Pt(15)
run_title.font.bold = True
run_title.font.color.rgb = COLOR_DARK_GREEN

# 4. Embedded Real Photo Table / Layout (Photo + Salutation side by side or centered)
portrait_path = r'C:\Users\inspy\.gemini\antigravity\scratch\medrobe-by-lene\assets\images\charlene_portrait.png'

# 2-Column Table for Photo and Opening Greeting
table = doc.add_table(rows=1, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = False

# Left Cell: Real Picture
cell_left = table.cell(0, 0)
cell_left.width = Inches(2.2)
p_img = cell_left.paragraphs[0]
p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_img = p_img.add_run()
run_img.add_picture(portrait_path, width=Inches(1.95))

# Photo Caption
p_caption = cell_left.add_paragraph()
p_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_caption.paragraph_format.space_before = Pt(3)
r_cap = p_caption.add_run('Charlene Odei Asare\nFinancial Secretary — Elect')
r_cap.font.name = 'Arial'
r_cap.font.size = Pt(8)
r_cap.font.bold = True
r_cap.font.color.rgb = COLOR_DARK_GREEN

# Right Cell: Opening Words
cell_right = table.cell(0, 1)
cell_right.width = Inches(4.5)
p_greet = cell_right.paragraphs[0]
p_greet.paragraph_format.line_spacing = 1.25
p_greet.paragraph_format.space_after = Pt(8)

r_sal = p_greet.add_run('Dear Medical Students & Esteemed FHUMSA Fraternity,\n\n')
r_sal.font.name = 'Georgia'
r_sal.font.size = Pt(10.5)
r_sal.font.bold = True
r_sal.font.color.rgb = COLOR_DARK

r_body1 = p_greet.add_run('To everyone who believed, encouraged, and voted—thank you. I am deeply humbled, honored, and grateful for the overwhelming trust you have placed in me. I won\'t take your trust for granted.\n\nWinning wasn\'t the finish line; it was the starting point. I step in with one mindset: WORK. SERVE. DELIVER.')
r_body1.font.name = 'Georgia'
r_body1.font.size = Pt(10.5)
r_body1.font.color.rgb = COLOR_DARK

# 5. Remaining Body Paragraphs below table
def add_body_p(text, bold_prefix=None, space_after=8, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(space_after)
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.font.name = 'Georgia'
        r_pre.font.size = Pt(10.5)
        r_pre.font.bold = True
        r_pre.font.color.rgb = COLOR_DARK
    r_body = p.add_run(text)
    r_body.font.name = 'Georgia'
    r_body.font.size = Pt(10.5)
    r_body.font.italic = italic
    r_body.font.color.rgb = COLOR_DARK
    return p

# Callout / Highlight Quote
p_quote = doc.add_paragraph()
p_quote.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_quote.paragraph_format.space_before = Pt(10)
p_quote.paragraph_format.space_after = Pt(10)
r_quote = p_quote.add_run('“I want my tenure felt, not just remembered. Seen, not just promised.”')
r_quote.font.name = 'Georgia'
r_quote.font.size = Pt(11.5)
r_quote.font.bold = True
r_quote.font.italic = True
r_quote.font.color.rgb = COLOR_PRIMARY

add_body_p('FHUMSA, you gave me the mandate; now watch me put it to work. Our commitment to transparency, student welfare, and prudent financial management begins today.')

# Conclusion Slogan
p_close = doc.add_paragraph()
p_close.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_close.paragraph_format.space_before = Pt(10)
p_close.paragraph_format.space_after = Pt(14)
r_close = p_close.add_run('CAMPAIGN OVER. SERVICE BEGINS.')
r_close.font.name = 'Arial'
r_close.font.size = Pt(11)
r_close.font.bold = True
r_close.font.color.rgb = COLOR_GOLD

# Sign-off
p_sign = doc.add_paragraph()
p_sign.paragraph_format.space_after = Pt(2)
r_sign1 = p_sign.add_run('Yours in Dedicated Service,\n\n')
r_sign1.font.name = 'Georgia'
r_sign1.font.size = Pt(10.5)
r_sign1.font.color.rgb = COLOR_DARK

r_sign_name = p_sign.add_run('Charlene Odei Asare\n')
r_sign_name.font.name = 'Georgia'
r_sign_name.font.size = Pt(11.5)
r_sign_name.font.bold = True
r_sign_name.font.color.rgb = COLOR_DARK_GREEN

r_sign_title = p_sign.add_run('Financial Secretary — Elect\nFamily Health University Medical Students\' Association (FHUMSA)\nFamily Health University College, Teshie, Accra, Ghana')
r_sign_title.font.name = 'Arial'
r_sign_title.font.size = Pt(9)
r_sign_title.font.color.rgb = COLOR_MUTED

# Footer Note (End of Press Release)
p_end = doc.add_paragraph()
p_end.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_end.paragraph_format.space_before = Pt(12)
r_end = p_end.add_run('###')
r_end.font.name = 'Arial'
r_end.font.bold = True
r_end.font.color.rgb = COLOR_MUTED

out_path = r'C:\Users\inspy\.gemini\antigravity\scratch\medrobe-by-lene\FHUMSA_Press_Release_Charlene_Odei_Asare.docx'
doc.save(out_path)
print(f'Successfully regenerated Word Document with real photo: {out_path}')
